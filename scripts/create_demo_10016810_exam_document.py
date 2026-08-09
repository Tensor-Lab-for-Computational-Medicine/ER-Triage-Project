from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "rubrics" / "demo_10016810_expected_physical_exam.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17232B"
MUTED = "607078"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CAUTION = "FFF4CE"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def font_run(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, *, bold=False, color=INK, size=11, after=6, before=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    font_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font_run(p.add_run(text), size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 7)
    run = p.add_run(text)
    font_run(run, size=16 if level == 1 else 13, bold=True, color=BLUE)
    return p


def add_callout(doc, label, text, fill=CAUTION):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    font_run(p.add_run(f"{label}: "), size=10.5, bold=True, color=DARK_BLUE)
    font_run(p.add_run(text), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_findings_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [1900, 5160, 2300]
    set_table_geometry(table, widths)
    headers = ["Maneuver", "Expected authored finding", "Teaching meaning"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font_run(p.add_run(text), size=9.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for maneuver, finding, meaning in rows:
        cells = table.add_row().cells
        for cell, width in zip(cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for index, text in enumerate((maneuver, finding, meaning)):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            font_run(p.add_run(text), size=9.25, bold=index == 0)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font_run(header.add_run("ER Triage Project  |  Faculty case specification"), size=8.5, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font_run(footer.add_run("demo_10016810  •  Clinician review required before summative use"), size=8.5, color=MUTED)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(4)
font_run(p.add_run("CLINICAL CASE SPECIFICATION"), size=10, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
font_run(p.add_run("Expected Physical Examination Findings"), size=24, bold=True, color=DARK_BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
font_run(p.add_run("demo_10016810  |  Acute appendicitis with generalized peritonitis"), size=12.5, color=MUTED)

meta = doc.add_table(rows=3, cols=2)
meta.style = "Table Grid"
set_table_geometry(meta, [2300, 7060])
metadata = [
    ("Patient", "66-year-old woman with severe abdominal pain"),
    ("Corrected triage", "Pain 8/10 • BP 98/48 mm Hg • SpO₂ 90% • ESI 2"),
    ("Teaching objective", "Recognize generalized peritonitis, stabilize physiologic threats, and escalate for urgent surgical care")
]
for row, (label, value) in zip(meta.rows, metadata):
    set_cell_shading(row.cells[0], LIGHT_GRAY)
    for cell, width in zip(row.cells, (2300, 7060)):
        set_cell_width(cell, width)
        set_cell_margins(cell)
    font_run(row.cells[0].paragraphs[0].add_run(label), size=9.5, bold=True, color=DARK_BLUE)
    font_run(row.cells[1].paragraphs[0].add_run(value), size=9.5)

add_callout(
    doc,
    "Core case pattern",
    "The examination should demonstrate a sick patient who avoids movement, has diffuse peritoneal irritation with maximal right-lower-quadrant findings, and has supportive appendiceal maneuvers. Generalized guarding, rigidity, rebound, and percussion tenderness are the most important findings."
)

add_heading(doc, "Expected findings by examination domain", 1)

add_heading(doc, "General and respiratory examination", 2)
general_rows = [
    ("General appearance", "Ill-appearing and uncomfortable; lies still to minimize abdominal pain; awake and speaking clearly.", "Movement avoidance supports peritoneal irritation; mental status remains intact."),
    ("Work of breathing", "Mildly increased work of breathing without severe respiratory distress.", "Prompts reassessment of SpO₂ 90% and response to oxygen."),
    ("Breath sounds", "Clear bilaterally without focal crackles or wheeze.", "Assesses thoracic alternatives and guides oxygen/fluid reassessment.")
]
add_findings_table(doc, general_rows)

add_heading(doc, "Abdominal examination", 2)
abdominal_rows = [
    ("Inspection", "Mild-to-moderate distention with shallow abdominal excursion; patient remains still because movement worsens pain.", "Supports ileus/peritoneal irritation and severe intra-abdominal disease."),
    ("Auscultation", "Hypoactive bowel sounds.", "May accompany ileus in generalized peritonitis; not diagnostic alone."),
    ("Light palpation", "Diffuse tenderness, maximal in the right lower quadrant at McBurney point.", "Localizes appendiceal inflammation within a generalized process."),
    ("Guarding / rigidity", "Involuntary guarding with generalized board-like rigidity, greatest in the right lower quadrant.", "High-urgency evidence of generalized peritoneal irritation."),
    ("Rebound", "Diffuse rebound tenderness, most pronounced in the right lower quadrant.", "Supports peritonitis; repeated painful maneuvers should be avoided."),
    ("Percussion", "Diffuse percussion tenderness, greatest in the right lower quadrant, with mild tympany over the distended abdomen.", "A less forceful way to demonstrate peritoneal irritation."),
    ("Rovsing sign", "Positive: left-lower-quadrant palpation elicits right-lower-quadrant pain.", "Supports appendiceal/peritoneal irritation but has limited sensitivity."),
    ("Psoas sign", "Positive: passive extension of the right hip increases right-lower-quadrant pain.", "Supports irritation near the iliopsoas, often with a retrocecal appendix."),
    ("Obturator sign", "Positive: internal rotation of the flexed right hip reproduces lower abdominal pain.", "Supports pelvic irritation; appendix position affects whether it is present."),
    ("Murphy sign", "Negative.", "Makes a biliary source less supported without excluding it by itself.")
]
add_findings_table(doc, abdominal_rows)

add_heading(doc, "How the simulator should present the findings", 1)
for text in (
    "Each finding should appear only after the learner selects the matching structured examination maneuver.",
    "The initial screen should show pain 8/10, not 0/10, and the patient should describe severe pain when asked.",
    "The examination record should use the catalog IDs for abdominal inspection, light palpation, guarding, rebound, bowel sounds, Rovsing, psoas, obturator, percussion, and Murphy sign.",
    "OpenEvidence should receive both the finding and its provenance, then explain its meaning, limitations, and effect on management.",
    "The local rubric should verify that the learner performed the maneuver; OpenEvidence should judge interpretation and clinical reasoning."
):
    add_bullet(doc, text)

add_heading(doc, "Clinical interpretation guardrails", 1)
add_callout(
    doc,
    "Faculty caution",
    "Rovsing, psoas, and obturator signs are supportive rather than required. Their sensitivities are low and their presence depends partly on appendix position. This authored case uses positive named signs to create a coherent teaching pattern, but learners should not be taught that all patients with appendicitis demonstrate all three.",
    fill=CAUTION
)
for text in (
    "Generalized involuntary guarding, rigidity, rebound, and percussion tenderness should drive urgency more strongly than any single named appendiceal sign.",
    "In an adult aged 66 years, atypical presentations and alternative diagnoses remain important; examination findings should guide but not replace definitive imaging and surgical evaluation.",
    "A normal or placeholder CT report must not override a clinically concerning examination. The current case report must remain explicitly labeled as templated until replaced with a clinician-reviewed diagnostic report."
):
    add_bullet(doc, text)

add_heading(doc, "Source basis", 1)
add_bullet(doc, "Podda M, et al. Diagnosis and Treatment of Acute Appendicitis: 2025 Edition of the World Society of Emergency Surgery Jerusalem Guidelines. JAMA Surgery. doi:10.1001/jamasurg.2025.6218")
add_bullet(doc, "Fugazzola P, et al. Guidelines for diagnosis and treatment of acute appendicitis in the elderly. World Journal of Emergency Surgery. 2020;15:19. doi:10.1186/s13017-020-00298-0")
add_bullet(doc, "Wagner JM, McKinney WP, Carpenter JL. Does This Patient Have Appendicitis? JAMA. 1996;276(19).")
add_text(doc, "Final case content and scoring should undergo clinician review and a complete faculty playthrough before summative use.", size=9.5, bold=True, color=RED, before=6)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
